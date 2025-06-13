from flask import Flask, request, jsonify ,Response,send_file
from flask_sqlalchemy import SQLAlchemy
from flask_bcrypt import Bcrypt
from flask_mail import Mail, Message
import random
import time
from flask_cors import CORS
from itsdangerous import URLSafeTimedSerializer as Serializer, SignatureExpired, BadSignature
import os
from dotenv import load_dotenv
import logging
from werkzeug.utils import secure_filename
from sentence_transformers import SentenceTransformer, util
import spacy
from datetime import datetime
from text_preprocess import clean_text
from utils import extract_text_from_pdf, extract_text_from_docx
import uuid
from fpdf import FPDF
import docx
from docx.shared import Pt
# Load environment variables
load_dotenv()

# Load NLP model
nlp = spacy.load(r'C:\Users\imaan\fyp\model-best\content\output\model-best')

app = Flask(__name__)
CORS(app, origins=["http://localhost:3000"])
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL')
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY')
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USE_SSL'] = True

db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
mail = Mail(app)

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Models
class UserSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, nullable=False)
    user_type = db.Column(db.String(20), nullable=False)
    session_id = db.Column(db.String(255), nullable=False)
    device_info = db.Column(db.Text, nullable=True)
    ip_address = db.Column(db.String(50), nullable=True)
    login_time = db.Column(db.DateTime, default=datetime.utcnow)
    last_activity = db.Column(db.DateTime, default=datetime.utcnow)
    is_active = db.Column(db.Boolean, default=True)

class HRUser(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    logged_out = db.Column(db.Boolean, default=False, nullable=False)

class CandidateUser(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    logged_out = db.Column(db.Boolean, default=False, nullable=False)

class Job(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    company_name = db.Column(db.String(255))
    title = db.Column(db.String(255))
    nature_of_role = db.Column(db.String(255))
    location = db.Column(db.String(255))
    education_requirement = db.Column(db.Text)
    work_experience = db.Column(db.Text)
    extra_skills = db.Column(db.Text)
    salary = db.Column(db.String(100))
    link = db.Column(db.String(255))
    recruiter_email = db.Column(db.String(255))
    job_requirements = db.Column(db.Text)
    skills_needed = db.Column(db.Text)
    industry = db.Column(db.String(255))
    deadline = db.Column(db.Date)
    hr_user_id = db.Column(db.Integer, db.ForeignKey('hr_user.id'))

class Application(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    candidate_id = db.Column(db.Integer, nullable=False)
    job_id = db.Column(db.Integer, nullable=False)
    cv_filename = db.Column(db.String(255), nullable=False)
    cv_text = db.Column(db.Text, nullable=False)
    report_filename = db.Column(db.String(255), nullable=True)

class ParsedCV(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    candidate_id = db.Column(db.Integer, db.ForeignKey('candidate_user.id'), nullable=False)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)
    parsed_data = db.Column(db.JSON, nullable=False)

    candidate = db.relationship('CandidateUser', backref=db.backref('parsed_cvs', lazy=True))
    job = db.relationship('Job', backref=db.backref('parsed_cvs', lazy=True))

class ApplicationStatus(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    candidate_id = db.Column(db.Integer, db.ForeignKey('candidate_user.id'), nullable=False)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)
    status = db.Column(db.String(50), default='under_review')  # under_review, selected, rejected
    feedback_sent = db.Column(db.Boolean, default=False)
    feedback_sent_at = db.Column(db.DateTime)
    last_updated = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    candidate = db.relationship('CandidateUser', backref=db.backref('application_statuses', lazy=True))
    job = db.relationship('Job', backref=db.backref('application_statuses', lazy=True))

with app.app_context():
    db.create_all()

# Temporary store for 2FA codes
two_factor_codes = {}

# Load SBERT model
sbert_model = SentenceTransformer('all-MiniLM-L6-v2')

# Label mapping for CV parsing
label_mapping = {
    'Job Title': ['JOB_TITLE'],
    'Location': ['LOCATION'],
    'Education Requirement': ['DEGREE', 'INSTITUTION', 'GRADES'],
    'Work Experience': ['EXPERIENCE', 'DURATION', 'PROJECT_TITLE', 'PROJECT_DESCRIPTION', 'CERTIFICATION', 'INTERNSHIP'],
    'Extra Skills': ['EXTRA_SKILLS', 'PROFICIENCY', 'LANGUAGE'],
    'Job Requirements': ['RESPONSIBILITIES', 'SUMMARY'],
    'Skills Needed': ['TECHNICAL_SKILL', 'SOFT_SKILL'],
    'Industry': ['EXPERIENCE']
}

@app.route('/')
def home():
    return jsonify({"message": "Welcome to the Smart-Hire API!"})

@app.route('/login', methods=['POST'])
def login():
    try:
        # Check if this is a resend OTP request
        if request.headers.get('X-Resend-OTP') == 'true':
            email = request.get_json().get('email')
            if not email:
                return jsonify({"message": "Email required for OTP resend"}), 400
                
            candidate = CandidateUser.query.filter_by(email=email).first()
            hr_user = HRUser.query.filter_by(email=email).first()
            user = candidate if candidate else hr_user
            
            if not user:
                return jsonify({"message": "User not found"}), 404
                
            return send_otp(user)

        # Original login logic
        data = request.get_json()
        email = data['email']
        password = data['password']

        candidate = CandidateUser.query.filter_by(email=email).first()
        if candidate and bcrypt.check_password_hash(candidate.password, password):
            return handle_user_login(candidate)

        hr_user = HRUser.query.filter_by(email=email).first()
        if hr_user and bcrypt.check_password_hash(hr_user.password, password):
            return handle_user_login(hr_user)

        return jsonify({"message": "Login failed. Please check your credentials."}), 401
    except Exception as e:
        logging.error(f"Error during login: {e}")
        return jsonify({"message": "An error occurred during login."}), 500
def handle_user_login(user):
    if user.id in two_factor_codes:
        del two_factor_codes[user.id]
    
    if user.logged_out:
        user.logged_out = False
        db.session.commit()
        return send_otp(user)
    
    # Create session for users who don't need OTP
    device_info = request.headers.get('User-Agent', 'unknown')
    ip_address = request.remote_addr
    session_id = str(uuid.uuid4())
    
    new_session = UserSession(
        user_id=user.id,
        user_type='candidate' if isinstance(user, CandidateUser) else 'hr',
        session_id=session_id,
        device_info=device_info,
        ip_address=ip_address
    )
    db.session.add(new_session)
    db.session.commit()
    
    return jsonify({
        "message": "Login successful",
        "user_id": user.id,
        "role": "candidate" if isinstance(user, CandidateUser) else "hr",
        "session_id": session_id,
        "skip_otp": True
    }), 200
def send_otp(user):
    try:
        otp = random.randint(100000, 999999)
        two_factor_codes[user.id] = {
            'code': otp,
            'timestamp': time.time()
        }

        msg = Message('Your OTP Code', sender=app.config['MAIL_USERNAME'], recipients=[user.email])
        msg.body = f"Your OTP code is: {otp}"
        mail.send(msg)

        return jsonify({
            "message": "OTP sent successfully",
            "user_id": user.id,
             "email": user.email,
            "role": "candidate" if isinstance(user, CandidateUser) else "hr"
        }), 200
    except Exception as e:
        logging.error(f"Error sending OTP: {e}")
        return jsonify({"message": "An error occurred while sending OTP."}), 500

@app.route('/verify', methods=['POST'])
def verify():
    try:
        data = request.get_json()
        user_id = int(data.get('user_id'))
        code = int(data.get('code'))

        candidate = CandidateUser.query.get(user_id)
        hr_user = HRUser.query.get(user_id)
        user = candidate if candidate else hr_user
        
        if not user:
            return jsonify({"error": "User not found"}), 404

        if user_id in two_factor_codes:
            stored_code = two_factor_codes[user_id]['code']
            timestamp = two_factor_codes[user_id]['timestamp']

            if time.time() - timestamp > 300:
                del two_factor_codes[user_id]
                return jsonify({"error": "OTP expired"}), 400

            if code == stored_code:
                del two_factor_codes[user_id]
                
                device_info = request.headers.get('User-Agent', 'unknown')
                ip_address = request.remote_addr
                session_id = str(uuid.uuid4())
                
                new_session = UserSession(
                    user_id=user_id,
                    user_type='candidate' if isinstance(user, CandidateUser) else 'hr',
                    session_id=session_id,
                    device_info=device_info,
                    ip_address=ip_address
                )
                db.session.add(new_session)
                db.session.commit()
                    
                return jsonify({
                    "message": "Verification successful",
                    "user_id": user_id,
                    "role": "candidate" if isinstance(user, CandidateUser) else "hr",
                    "session_id": session_id
                }), 200
            else:
                return jsonify({"error": "Invalid OTP"}), 400

        return jsonify({"error": "No OTP found. Please request a new OTP."}), 404
    except Exception as e:
        logging.error(f"Error during verification: {e}")
        return jsonify({"message": "An error occurred during verification."}), 500
from datetime import timezone  # Add timezone to imports

@app.route('/api/logout', methods=['POST'])
def logout():
    try:
        data = request.get_json()
        user_id = data.get('user_id')
        user_type = data.get('user_type')  # This should be 'candidate' or 'hr'
        
        logging.info(f"Logout request for user_id: {user_id}, type: {user_type}")
        
        if not user_id:
            return jsonify({"message": "User ID required", "success": False}), 400

        # Validate user_type
        if user_type not in ['candidate', 'hr']:
            return jsonify({"message": "Invalid user type", "success": False}), 400

        # Get the correct user model based on the provided user_type
        user = None
        if user_type == 'candidate':
            user = db.session.get(CandidateUser, user_id)
        elif user_type == 'hr':
            user = db.session.get(HRUser, user_id)

        if not user:
            return jsonify({"message": "User not found", "success": False}), 404

        logging.info(f"Logging out {user_type} user {user_id}")

        # Update logged_out status
        user.logged_out = True
        
        # Deactivate all sessions
        active_sessions = db.session.query(UserSession).filter(
            UserSession.user_id == user_id, 
            UserSession.is_active == True
        ).all()
        
        for session in active_sessions:
            session.is_active = False
            session.logout_time = datetime.now(timezone.utc)
        
        # Clear any 2FA codes
        if user_id in two_factor_codes:
            del two_factor_codes[user_id]

        db.session.commit()

        logging.info(f"Successfully logged out {user_type} user {user_id}")
        
        return jsonify({
            "message": f"{user_type.capitalize()} user logged out successfully",
            "success": True,
            "user_type": user_type  # Return the same user_type we received
        }), 200
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Logout error: {str(e)}", exc_info=True)
        return jsonify({
            "message": "Logout failed. Please try again.",
            "success": False
        }), 500
@app.route('/check-logout-status/<int:user_id>', methods=['GET'])
def check_logout_status(user_id):
    try:
        candidate = CandidateUser.query.get(user_id)
        hr_user = HRUser.query.get(user_id)
        user = candidate if candidate else hr_user
        
        if not user:
            return jsonify({"message": "User not found"}), 404
            
        return jsonify({"logged_out": user.logged_out}), 200
    except Exception as e:
        logging.error(f"Error checking logout status: {e}")
        return jsonify({"message": "An error occurred"}), 500

@app.route('/hr/<int:user_id>', methods=['GET'])
def get_hr(user_id):
    try:
        hr = HRUser.query.get(user_id)
        if hr:
            return jsonify({
                'name': hr.name,
                'email': hr.email
            }), 200
        return jsonify({"message": "HR not found"}), 404
    except Exception as e:
        logging.error(f"Error fetching HR data: {e}")
        return jsonify({"message": "An error occurred"}), 500

@app.route('/jobs', methods=['GET'])
def list_jobs():
    try:
        jobs = Job.query.all()
        job_list = [
            {
                "id": job.id,
                "company_name": job.company_name,
                "job_title": job.title,
                "nature_of_role": job.nature_of_role,
                "location": job.location,
                "education_requirement": job.education_requirement,
                "work_experience": job.work_experience,
                "extra_skills": job.extra_skills,
                "salary": job.salary,
                "link": job.link,
                "recruiter_email": job.recruiter_email,
                "job_requirements": job.job_requirements,
                "skills_needed": job.skills_needed,
                "industry": job.industry,
                "deadline": job.deadline.strftime('%Y-%m-%d') if isinstance(job.deadline, datetime) else job.deadline
            } for job in jobs
        ]
        return jsonify(job_list), 200
    except Exception as e:
        logging.error(f"Error fetching job listings: {e}")
        return jsonify({"message": "An error occurred while fetching job listings."}), 500

@app.route('/hr/post-job', methods=['POST'])
def post_job():
    try:
        data = request.get_json()
        deadline = datetime.strptime(data['deadline'], '%Y-%m-%d').date()
        new_job = Job(
            company_name=data['company_name'],
            title=data['job_title'],
            nature_of_role=data['nature_of_role'],
            location=data['location'],
            education_requirement=data['education_requirement'],
            work_experience=data['work_experience'],
            extra_skills=data['extra_skills'],
            salary=data['salary'],
            link=data['link'],
            recruiter_email=data['recruiter_email'],
            job_requirements=data['job_requirements'],
            skills_needed=data['skills_needed'],
            industry=data['industry'],
            deadline=deadline,
            hr_user_id=data['hr_user_id']
        )

        db.session.add(new_job)
        db.session.commit()
        return jsonify({"message": "Job posted successfully!"}), 201
    except Exception as e:
        logging.error(f"Error posting job: {e}")
        return jsonify({"message": "An error occurred"}), 500

@app.route('/hr/edit-job/<int:job_id>', methods=['PUT'])
def edit_job(job_id):
    try:
        data = request.get_json()
        job = Job.query.get(job_id)
        if not job:
            return jsonify({"message": "Job not found"}), 404

        job.company_name = data['company_name']
        job.title = data['job_title']
        job.nature_of_role = data['nature_of_role']
        job.location = data['location']
        job.education_requirement = data['education_requirement']
        job.work_experience = data['work_experience']
        job.extra_skills = data['extra_skills']
        job.salary = data['salary']
        job.link = data['link']
        job.recruiter_email = data['recruiter_email']
        job.job_requirements = data['job_requirements']
        job.skills_needed = data['skills_needed']
        job.industry = data['industry']
        job.deadline = datetime.strptime(data['deadline'], '%Y-%m-%d').date()
        
        db.session.commit()
        return jsonify({"message": "Job updated successfully!"}), 200
    except Exception as e:
        logging.error(f"Error editing job: {e}")
        return jsonify({"message": "An error occurred"}), 500

@app.route('/hr/delete-job/<int:job_id>', methods=['DELETE'])
def delete_job(job_id):
    try:
        job = Job.query.get(job_id)
        if not job:
            return jsonify({"message": "Job not found"}), 404

        # Delete related application statuses
        application_statuses = ApplicationStatus.query.filter_by(job_id=job_id).all()
        for status in application_statuses:
            db.session.delete(status)

        # Delete related applications
        applications = Application.query.filter_by(job_id=job_id).all()
        for app in applications:
            db.session.delete(app)

        # Delete related parsed CVs
        parsed_cvs = ParsedCV.query.filter_by(job_id=job_id).all()
        for cv in parsed_cvs:
            db.session.delete(cv)

        # Delete the job
        db.session.delete(job)
        db.session.commit()

        return jsonify({"message": "Job and all associated data deleted successfully!"}), 200
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error deleting job: {e}")
        return jsonify({"message": "An error occurred while deleting the job"}), 500

@app.route('/candidate/<int:user_id>', methods=['GET'])
def get_candidate(user_id):
    try:
        candidate = CandidateUser.query.get(user_id)
        if candidate:
            return jsonify({
                'name': candidate.name,
                'email': candidate.email
            }), 200
        return jsonify({"message": "Candidate not found"}), 404
    except Exception as e:
        logging.error(f"Error fetching candidate data: {e}")
        return jsonify({"message": "An error occurred"}), 500

@app.route('/candidate/signup', methods=['POST'])
def candidate_signup():
    try:
        data = request.get_json()
        name = data['name']
        email = data['email']
        password = data['password']

        if CandidateUser.query.filter_by(email=email).first():
            return jsonify({"message": "Candidate already exists!"}), 400

        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        new_candidate = CandidateUser(name=name, email=email, password=hashed_password)
        db.session.add(new_candidate)
        db.session.commit()

        return jsonify({"message": "Candidate signup successful!"}), 201
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error during candidate signup: {e}")
        return jsonify({"message": "An error occurred during signup."}), 500

@app.route('/candidate/applied-jobs/<int:user_id>', methods=['GET'])
def get_applied_jobs(user_id):
    try:
        applications = Application.query.filter_by(candidate_id=user_id).all()
        job_ids = [application.job_id for application in applications]
        jobs = Job.query.filter(Job.id.in_(job_ids)).all()

        job_list = [{
            'id': job.id,
            'company_name': job.company_name,
            'job_title': job.title,
            'nature_of_role': job.nature_of_role,
            'location': job.location,
            'education_requirement': job.education_requirement,
            'work_experience': job.work_experience,
            'extra_skills': job.extra_skills,
            'salary': job.salary,
            'link': job.link,
            'recruiter_email': job.recruiter_email,
            'job_requirements': job.job_requirements,
            'skills_needed': job.skills_needed,
            'industry': job.industry,
            'deadline': job.deadline.strftime('%Y-%m-%d') if isinstance(job.deadline, datetime) else job.deadline
        } for job in jobs]

        return jsonify(job_list), 200
    except Exception as e:
        logging.error(f"Error fetching applied jobs for user {user_id}: {e}")
        return jsonify({"message": "An error occurred while fetching applied jobs."}), 500

@app.route('/candidate/apply-job', methods=['POST'])
def apply_job():
    try:
        if 'cv' not in request.files:
            return jsonify({'error': 'No file part'}), 400

        file = request.files['cv']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        if not allowed_file(file.filename):
            return jsonify({'error': 'Unsupported file format'}), 400

        cv_filename = secure_filename(file.filename)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], cv_filename))
        report_filename = f"report_{cv_filename}.pdf"  # Generate report filename

        if cv_filename.endswith('.pdf'):
            cv_text = extract_text_from_pdf(file)
        elif cv_filename.endswith('.docx'):
            cv_text = extract_text_from_docx(file)
        else:
            return jsonify({'error': 'Unsupported file format'}), 400

        cv_text = clean_text(cv_text)

        candidate_id = request.form.get('candidate_id')
        job_id = request.form.get('job_id')

        if not candidate_id or not job_id:
            return jsonify({'error': 'Missing candidate_id or job_id'}), 400

        new_application = Application(
            candidate_id=candidate_id,
            job_id=job_id,
            cv_filename=cv_filename,
            report_filename=report_filename,  # Add this
            cv_text=cv_text
        )
        db.session.add(new_application)
        
        # Create initial application status
        app_status = ApplicationStatus(
            candidate_id=candidate_id,
            job_id=job_id,
            status='under_review'
        )
        db.session.add(app_status)
        
        db.session.commit()

        extract_entities(candidate_id, job_id, cv_text)

        return jsonify({'message': 'Application submitted successfully'}), 201
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error applying for job: {e}")
        return jsonify({'error': str(e)}), 500
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']    
def extract_entities(candidate_id, job_id, cv_text):
    try:
        doc = nlp(cv_text)

        parsed_data = {
            'name': [ent.text for ent in doc.ents if ent.label_ == 'NAME'],
            'email': [ent.text for ent in doc.ents if ent.label_ == 'EMAIL'],
            'phone_number': [ent.text for ent in doc.ents if ent.label_ == 'PHONE_NUMBER'],
            'location': [ent.text for ent in doc.ents if ent.label_ == 'LOCATION'],
            'technical_skills': [ent.text for ent in doc.ents if ent.label_ == 'TECHNICAL_SKILL'],
            'soft_skills': [ent.text for ent in doc.ents if ent.label_ == 'SOFT_SKILL'],
            'job_title': [ent.text for ent in doc.ents if ent.label_ == 'JOB_TITLE'],
            'company': [ent.text for ent in doc.ents if ent.label_ == 'COMPANY'],
            'duration': [ent.text for ent in doc.ents if ent.label_ == 'DURATION'],
            'responsibilities': [ent.text for ent in doc.ents if ent.label_ == 'RESPONSIBILITIES'],
            'degree': [ent.text for ent in doc.ents if ent.label_ == 'DEGREE'],
            'institution': [ent.text for ent in doc.ents if ent.label_ == 'INSTITUTION'],
            'grades': [ent.text for ent in doc.ents if ent.label_ == 'GRADES'],
            'project_title': [ent.text for ent in doc.ents if ent.label_ == 'PROJECT_TITLE'],
            'project_description': [ent.text for ent in doc.ents if ent.label_ == 'PROJECT_DESCRIPTION'],
            'project_link': [ent.text for ent in doc.ents if ent.label_ == 'PROJECT_LINK'],
            'language': [ent.text for ent in doc.ents if ent.label_ == 'LANGUAGE'],
            'proficiency': [ent.text for ent in doc.ents if ent.label_ == 'PROFICIENCY'],
            'certification': [ent.text for ent in doc.ents if ent.label_ == 'CERTIFICATION'],
            'extra_skills': [ent.text for ent in doc.ents if ent.label_ == 'EXTRA_SKILLS'],
            'summary': [ent.text for ent in doc.ents if ent.label_ == 'SUMMARY'],
            'experience': [ent.text for ent in doc.ents if ent.label_ == 'EXPERIENCE'],
            'internship': [ent.text for ent in doc.ents if ent.label_ == 'INTERNSHIP'],
            'programming_language': [ent.text for ent in doc.ents if ent.label_ == 'PROGRAMMING_LANGUAGE']
        }

        parsed_cv = ParsedCV(
            candidate_id=candidate_id,
            job_id=job_id,
            parsed_data=parsed_data
        )
        db.session.add(parsed_cv)
        db.session.commit()

    except Exception as e:
        logging.error(f"Error during entity extraction: {e}")

@app.route('/forgot-password', methods=['POST'])
def forgot_password():
    try:
        data = request.get_json()
        email = data['email']
        
        candidate = CandidateUser.query.filter_by(email=email).first()
        hr_user = HRUser.query.filter_by(email=email).first()

        if not candidate and not hr_user:
            return jsonify({"message": "User not found"}), 404

        user = candidate if candidate else hr_user
        user_role = 'candidate' if candidate else 'hr'

        s = Serializer(app.config['SECRET_KEY'])
        token = s.dumps({'user_id': user.id, 'role': user_role}, salt='password-reset')

        base_url = os.getenv('FRONTEND_URL', 'http://localhost:3000')
        reset_link = f"{base_url}/reset-password?token={token}&user={user_role}"
        msg = Message('Password Reset Request', sender=app.config['MAIL_USERNAME'], recipients=[email])
        msg.body = f"Click on the following link to reset your password: {reset_link}"
        mail.send(msg)

        return jsonify({"message": "Password reset link sent successfully"}), 200
    except Exception as e:
        logging.error(f"Error during password reset request: {e}")
        return jsonify({"message": "An error occurred during password reset request."}), 500

@app.route('/validate-token/<token>', methods=['GET'])
def validate_token(token):
    try:
        s = Serializer(app.config['SECRET_KEY'])
        try:
            data = s.loads(token, salt='password-reset', max_age=3600)
            return jsonify({"valid": True}), 200
        except SignatureExpired:
            return jsonify({"valid": False, "message": "The token has expired"}), 400
        except BadSignature:
            return jsonify({"valid": False, "message": "Invalid token"}), 400
    except Exception as e:
        logging.error(f"Error during token validation: {e}")
        return jsonify({"message": "An error occurred during token validation."}), 500

@app.route('/reset-password', methods=['POST'])
def reset_password():
    try:
        data = request.get_json()
        token = data['token']
        new_password = data['newPassword']

        s = Serializer(app.config['SECRET_KEY'])
        try:
            user_data = s.loads(token, salt='password-reset', max_age=3600)
            user_id = user_data['user_id']
            user_role = user_data.get('role', None)

            if user_role not in ['candidate', 'hr']:
                return jsonify({"message": "Invalid user role"}), 400

            if user_role == 'candidate':
                user = CandidateUser.query.get(user_id)
            else:
                user = HRUser.query.get(user_id)

            if not user:
                return jsonify({"message": "User not found"}), 404

            hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
            user.password = hashed_password
            db.session.commit()

            return jsonify({"message": "Password has been successfully reset"}), 200
        except SignatureExpired:
            return jsonify({"message": "The token has expired"}), 400
        except BadSignature:
            return jsonify({"message": "Invalid token"}), 400
    except Exception as e:
        logging.error(f"Error during password reset: {e}")
        return jsonify({"message": "An error occurred during password reset."}), 500

@app.route('/contact-us', methods=['POST'])
def contact_us():
    data = request.json
    name = data.get('name')
    email = data.get('email')
    message = data.get('message')

    if not name or not email or not message:
        return jsonify({'message': 'All fields are required.'}), 400

    try:
        msg = Message(
            subject=f"Query from {name}",
            recipients=[os.getenv('MAIL_USERNAME')],
            body=f"Name: {name}\nEmail: {email}\n\nMessage:\n{message}",
            sender=os.getenv('MAIL_DEFAULT_SENDER')
        )
        mail.send(msg)
        return jsonify({'message': 'Message sent successfully!'}), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'message': 'Failed to send message.'}), 500

@app.route('/hr/job-candidates/<int:job_id>', methods=['GET'])
def get_job_applicants_with_scores(job_id):
    try:
        job = Job.query.get(job_id)
        if not job:
            return jsonify({"message": "Job not found"}), 404

        applications = Application.query.filter_by(job_id=job_id).all()
        candidate_ids = [app.candidate_id for app in applications]

        parsed_cvs = ParsedCV.query.filter_by(job_id=job_id).all()
        parsed_cv_dict = {
            cv.candidate_id: {
                "cv_id": cv.id,
                "parsed_data": cv.parsed_data
            }
            for cv in parsed_cvs
        }

        job_data = job.__dict__
        job_text = ' '.join([
            str(job_data.get('education_requirement', '')),
            str(job_data.get('work_experience', '')),
            str(job_data.get('extra_skills', '')),
            str(job_data.get('job_requirements', '')),
            str(job_data.get('skills_needed', ''))
        ])
        job_text_clean = clean_text(job_text)

        results = []
        for candidate_id in candidate_ids:
            candidate = CandidateUser.query.get(candidate_id)
            parsed_cv_info = parsed_cv_dict.get(candidate_id)

            if not parsed_cv_info:
                sbert_score = 0.0
                parsed_data = None
                cv_id = None
            else:
                parsed_data = parsed_cv_info["parsed_data"]
                cv_id = parsed_cv_info["cv_id"]

                cv_text = ' '.join([
                    ' '.join(parsed_data.get(label.lower(), []))
                    for labels in label_mapping.values()
                    for label in labels
                ])
                cv_text_clean = clean_text(cv_text)

                sbert_score = sbert_similarity_score(job_text_clean, cv_text_clean)

            results.append({
                "candidate_id": candidate.id,
                "candidate_name": candidate.name,
                "candidate_email": candidate.email,
                "cv_id": cv_id,
                "sbert_score": round(float(sbert_score) * 100, 2),
                "parsed_cv": parsed_data
            })

        return jsonify(results), 200

    except Exception as e:
        logging.error(f"Error fetching candidates for job {job_id}: {e}")
        return jsonify({"message": f"An error occurred: {str(e)}"}), 500

def sbert_similarity_score(text1, text2):
    if not text1 or not text2:
        return 0.0
    try:
        embeddings = sbert_model.encode([text1, text2], convert_to_tensor=True)
        return float(util.pytorch_cos_sim(embeddings[0], embeddings[1])[0])
    except Exception as e:
        logging.error(f"Error during SBERT similarity: {e}")
        return 0.0

import re
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS

def sbert_similarity_score(text1, text2):
    if not text1 or not text2:
        return 0.0
    try:
        embeddings = sbert_model.encode([text1, text2], convert_to_tensor=True)
        return float(util.pytorch_cos_sim(embeddings[0], embeddings[1])[0])
    except Exception as e:
        logging.error(f"Error during SBERT similarity: {e}")
        return 0.0

def extract_keywords(text):
    if not text:
        return []
    words = re.findall(r'\b[a-zA-Z0-9.#\+]+\b', text.lower())
    keywords = [word for word in words if word not in ENGLISH_STOP_WORDS and len(word) > 2]
    return list(set(keywords))

def keyword_match_score(text, keywords):
    matches = [kw for kw in keywords if kw.lower() in text.lower()]
    return len(matches) / len(keywords) * 100 if keywords else 0.0

@app.route('/api/sbert-feedback/<int:job_id>/<int:cv_id>', methods=['GET'])
def sbert_feedback(job_id, cv_id):
    job = Job.query.get(job_id)
    parsed_cv = ParsedCV.query.get(cv_id)

    if not job:
        return jsonify({"message": f"Job with ID {job_id} not found"}), 404
    if not parsed_cv:
        return jsonify({"message": f"Parsed CV with ID {cv_id} not found"}), 404

    parsed = parsed_cv.parsed_data

    field_mapping = {
        'Job Requirements': ['summary', 'responsibilities'],
        'Skills Needed': ['technical_skills', 'soft_skills', 'extra_skills', 'programming_language'],
        'Industry': ['experience', 'job_title'],
        'Education Requirement': ['degree', 'institution', 'grades'],
        'Work Experience': ['experience', 'duration', 'certification', 'project_title', 'project_description', 'internship'],
        'Extra Skills': ['extra_skills', 'proficiency', 'language', 'programming_language']
    }

    job_fields = {
        'Job Requirements': job.job_requirements,
        'Skills Needed': job.skills_needed,
        'Industry': job.industry,
        'Education Requirement': job.education_requirement,
        'Work Experience': job.work_experience,
        'Extra Skills': job.extra_skills
    }

    matching_items = []
    missing_items = []
    threshold = 0.45

    for job_label, job_value in job_fields.items():
        if not job_value:
            continue

        candidate_text = ' '.join([
            ' '.join(parsed.get(label, []))
            for label in field_mapping.get(job_label, [])
        ])

        # Special logic for Work Experience
        if job_label == 'Work Experience':
            durations = parsed.get('duration', [])
            titles = parsed.get('job_title', [])
            internships = parsed.get('internship', [])
            projects = parsed.get('project_title', [])

            work_history = []

            # Build work experience list using available data
            for i in range(len(durations)):
                work_type = 'job'
                if i < len(internships):
                    work_type = 'internship'
                elif i < len(projects):
                    work_type = 'project'

                work_history.append({
                    'type': work_type,
                    'duration': durations[i]
                })

            total_months = calculate_total_experience(work_history, job_value)
            min_required_years = extract_min_years_requirement(job_value)
            min_required_months = min_required_years * 12
            exp_score = min(total_months / min_required_months, 1.0) * 100 if min_required_months else 100
            exp_score = round(exp_score, 2)

            item_data = {
                'field': job_label,
                'job_data': job_value,
                'cv_data': f"{round(total_months, 1)} months of experience",
                'score': exp_score
            }

            if exp_score >= threshold * 100:
                matching_items.append(item_data)
            else:
                missing_items.append(item_data)

            continue  # Skip standard SBERT matching for this field

        # Default SBERT + keyword logic for other fields
        total_lines = job_value.strip().split('\n')
        combined_score = 0
        total_lines_count = 0

        for line in total_lines:
            line = line.strip()
            if not line:
                continue

            semantic_score = sbert_similarity_score(line, candidate_text) * 100
            keywords = extract_keywords(line)
            keyword_score = keyword_match_score(candidate_text, keywords)

            final_score = 0.7 * semantic_score + 0.3 * keyword_score
            if semantic_score < 30 and keyword_score == 0:
                final_score = 0

            combined_score += final_score
            total_lines_count += 1

        avg_score = round(combined_score / total_lines_count, 2) if total_lines_count else 0

        item_data = {
            'field': job_label,
            'job_data': job_value,
            'cv_data': candidate_text,
            'score': avg_score
        }

        if avg_score >= threshold * 100:
            matching_items.append(item_data)
        else:
            missing_items.append(item_data)

    # Overall SBERT score
    job_text = ' '.join([str(val or '') for val in job_fields.values()])
    candidate_text = ' '.join([
        ' '.join(parsed.get(field, []))
        for fields in field_mapping.values() for field in fields
    ])

    semantic_score = sbert_similarity_score(job_text, candidate_text) * 100
    keywords = extract_keywords(job_text)
    keyword_score = keyword_match_score(candidate_text, keywords)
    sbert_score = round(0.7 * semantic_score + 0.3 * keyword_score, 2)

    if sbert_score >= 85:
        feedback = "🌟 Excellent match! Your profile aligns perfectly with the job requirements."
    elif sbert_score >= 65:
        feedback = "👍 Good match. You are close to meeting most job requirements. A few improvements could help."
    elif sbert_score >= 45:
        feedback = "🧐 Partial match. You have some relevant skills, but need to align more with the job description."
    else:
        feedback = "❌ Not a good match currently. Consider gaining more relevant experience or skills for this role."

    return jsonify({
        "sbert_score": sbert_score,
        "feedback_message": feedback,
        "matching_items": matching_items,
        "missing_items": missing_items
    }), 200
# Helper Functions
def parse_duration_to_months(duration_text):
    """Convert duration text to months"""
    if not duration_text:
        return 0
    
    month_map = {m.lower(): i+1 for i, m in enumerate([
        'January', 'February', 'March', 'April', 'May', 'June',
        'July', 'August', 'September', 'October', 'November', 'December'
    ])}
    
    try:
        # Try "June 2023 - August 2023" or "June 2023 to August 2023" format
        parts = re.split(r'\s*-\s*|\s+to\s+', duration_text.strip(), flags=re.I)
        if len(parts) >= 2:
            # Handle month-year formats
            if all(any(c.isalpha() for c in part) for part in parts[:2]):  # Contains letters (month names)
                start_parts = parts[0].split()
                end_parts = parts[1].split()
                
                if len(start_parts) >= 2 and len(end_parts) >= 2:
                    start_month = start_parts[0]
                    start_year = start_parts[-1]  # Get last part as year
                    end_month = end_parts[0]
                    end_year = end_parts[-1]
                    
                    start = int(start_year) * 12 + month_map.get(start_month.lower(), 0)
                    end = int(end_year) * 12 + month_map.get(end_month.lower(), 0)
                    return max(end - start, 1)  # Minimum 1 month
            
            # Handle MM/YYYY formats
            if '/' in parts[0] and '/' in parts[1]:
                start_month, start_year = parts[0].split('/')[:2]
                end_month, end_year = parts[1].split('/')[:2]
                
                start = int(start_year) * 12 + int(start_month)
                end = int(end_year) * 12 + int(end_month)
                return max(end - start, 1)
                
        # Try "3 months" or "2 years" format
        match = re.match(r'(\d+)\s*(month|year)s?', duration_text, re.I)
        if match:
            num = int(match.group(1))
            unit = match.group(2).lower()
            return num if unit == 'month' else num * 12
            
    except Exception as e:
        logging.error(f"Error parsing duration '{duration_text}': {e}")
    
    # Default to 1 month if we can't parse but there's some duration text
    return 1 if duration_text.strip() else 0

def calculate_total_experience(work_history, job_requirements):
    """Calculate total relevant experience in months with improved logic"""
    if not work_history:
        return 0
        
    job_req_text = (job_requirements or '').lower()
    
    # Determine accepted experience types
    accepts_internships = any(x in job_req_text for x in ['internship', 'intern'])
    accepts_projects = 'project' in job_req_text
    strict_mode = 'only work experience' in job_req_text
    
    total_months = 0
    for exp in work_history:
        if not isinstance(exp, dict):
            continue
            
        duration = parse_duration_to_months(exp.get('duration', ''))
        
        # Job experience (full weight)
        if exp.get('type') == 'job':
            total_months += duration
        
        # Internships (reduced weight if accepted)
        elif exp.get('type') == 'internship' and (accepts_internships or not strict_mode):
            total_months += duration * 0.7  # 70% weight
        
        # Projects (further reduced weight if accepted)
        elif exp.get('type') == 'project' and accepts_projects:
            total_months += duration * 0.5  # 50% weight
    
    return total_months

def extract_min_years_requirement(text):
    """Extract minimum years requirement from text"""
    if not text:
        return 0
        
    match = re.search(r'(\d+)\s*-\s*(\d+)\s*years', text, re.I)
    if match:
        return int(match.group(1))
        
    match = re.search(r'at least\s*(\d+)\s*years', text, re.I)
    if match:
        return int(match.group(1))
        
    match = re.search(r'minimum\s*(\d+)\s*years', text, re.I)
    if match:
        return int(match.group(1))
        
    match = re.search(r'(\d+)\s*\+\s*years', text, re.I)
    if match:
        return int(match.group(1))
        
    return 0

@app.route('/candidate/application-status/<int:candidate_id>', methods=['GET'])
def get_application_status(candidate_id):
    try:
        statuses = ApplicationStatus.query.filter_by(candidate_id=candidate_id).all()
        
        if not statuses:
            return jsonify({"message": "No applications found"}), 404
            
        result = []
        for status in statuses:
            job = Job.query.get(status.job_id)
            if not job:
                continue
                
            result.append({
                "job_id": status.job_id,
                "job_title": job.title,
                "company_name": job.company_name,
                "status": status.status,
                "feedback_sent": status.feedback_sent,
                "feedback_sent_at": status.feedback_sent_at.isoformat() if status.feedback_sent_at else None,
                "last_updated": status.last_updated.isoformat()
            })
        
        return jsonify(result), 200
        
    except Exception as e:
        logging.error(f"Error fetching application status for candidate {candidate_id}: {e}")
        return jsonify({"message": "An error occurred while fetching application status"}), 500
# Add these imports at the top if not already present
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def send_candidate_notification(job_id, candidate_id, status, feedback=None):
    try:
        candidate = CandidateUser.query.get(candidate_id)
        job = Job.query.get(job_id)
        
        if not candidate or not job:
            logging.error(f"Candidate {candidate_id} or job {job_id} not found")
            return False

        # Get parsed CV data for feedback details
        parsed_cv = ParsedCV.query.filter_by(
            job_id=job_id,
            candidate_id=candidate_id
        ).first()
        
        # Get SBERT feedback if available
        feedback_details = None
        if parsed_cv:
            try:
                feedback_response = sbert_feedback(job_id, parsed_cv.id)
                if feedback_response[1] == 200:
                    feedback_details = feedback_response[0].json
            except Exception as e:
                logging.error(f"Error getting feedback details: {e}")

        # Update application status
        app_status = ApplicationStatus.query.filter_by(
            candidate_id=candidate_id,
            job_id=job_id
        ).first()
        
        if not app_status:
            app_status = ApplicationStatus(
                candidate_id=candidate_id,
                job_id=job_id,
                status=status
            )
            db.session.add(app_status)
        
        app_status.status = status
        app_status.feedback_sent = True
        app_status.feedback_sent_at = datetime.utcnow()
        db.session.commit()

        # Prepare email content
        subject = f"Application Update for {job.title} at {job.company_name}"

        if status == 'selected':
            body = f"""<html>
            <body style="font-family: Arial, sans-serif; color: #000000; line-height: 1.6;">
                <p>Dear {candidate.name},</p>
                
                <p>Congratulations! We're pleased to inform you that your application for the <strong>{job.title}</strong> 
                position at <strong>{job.company_name}</strong> has been successful based on our initial screening.</p>
            """
        else:
            body = f"""<html>
            <body style="font-family: Arial, sans-serif; color: #000000; line-height: 1.6;">
                <p>Dear {candidate.name},</p>
                
                <p>Thank you for applying for the <strong>{job.title}</strong> position at <strong>{job.company_name}</strong>. 
                After careful consideration, we regret to inform you that your application 
                was not successful at this time.</p>
            """

        # Add detailed matching/missing items
        if feedback_details:
            if feedback_details.get('matching_items'):
                body += """
                <h3 style="color: #000000; margin-top: 20px;">Your application matched well in these areas:</h3>
                <ul style="padding-left: 20px;">
                """
                for item in feedback_details['matching_items']:
                    body += f"""
                    <li style="margin-bottom: 15px;">
                        <strong>{item['field']}:</strong><br>
                        <span style="color: #333333;">Job Requirement: {item['job_data']}</span><br>
                        <span style="color: #333333;">Your Experience: {item['cv_data']}</span>
                    </li>
                    """
                body += "</ul>"
            
            if feedback_details.get('missing_items'):
                body += """
                <h3 style="color: #000000; margin-top: 20px;">Areas where your application could be strengthened:</h3>
                <ul style="padding-left: 20px;">
                """
                for item in feedback_details['missing_items']:
                    body += f"""
                    <li style="margin-bottom: 15px;">
                        <strong>{item['field']}:</strong><br>
                        <span style="color: #333333;">Job Requirement: {item['job_data']}</span><br>
                        <span style="color: #333333;">Your Experience: {item['cv_data'] if item['cv_data'] else '(Not found in your application)'}</span>
                    </li>
                    """
                body += "</ul>"

        # Add general feedback if available
        if feedback:
            body += f"""
            <h3 style="color: #000000; margin-top: 20px;">Additional Feedback:</h3>
            <p style="color: #333333;">{feedback}</p>
            """

        body += f"""
        <p style="margin-top: 30px;">
            Best regards,<br>
            <strong>{job.company_name} Recruitment Team</strong>
        </p>
        </body>
        </html>
        """

        msg = Message(
            subject=subject,
            sender=app.config['MAIL_USERNAME'],
            recipients=[candidate.email]
        )
        msg.html = body
        mail.send(msg)
        return True
    except Exception as e:
        logging.error(f"Error sending notification to candidate {candidate_id}: {e}")
        db.session.rollback()
        return False
     
@app.route('/candidate/application-status/<int:candidate_id>/<int:job_id>', methods=['GET'])
def get_single_application_status(candidate_id, job_id):
    try:
        status = ApplicationStatus.query.filter_by(
            candidate_id=candidate_id,
            job_id=job_id
        ).first()
        
        if not status:
            return jsonify({"message": "Application not found"}), 404
            
        job = Job.query.get(job_id)
        if not job:
            return jsonify({"message": "Job not found"}), 404
            
        return jsonify({
            "job_id": status.job_id,
            "job_title": job.title,
            "company_name": job.company_name,
            "status": status.status,
            "feedback_sent": status.feedback_sent,
            "feedback_sent_at": status.feedback_sent_at.isoformat() if status.feedback_sent_at else None,
            "last_updated": status.last_updated.isoformat()
        }), 200
        
    except Exception as e:
        logging.error(f"Error fetching application status for candidate {candidate_id}, job {job_id}: {e}")
        return jsonify({"message": "An error occurred while fetching application status"}), 500
@app.route('/hr/auto-evaluate-candidates/<int:job_id>', methods=['POST'])
def auto_evaluate_candidates(job_id):
    try:
        job = Job.query.get(job_id)
        if not job:
            return jsonify({"message": "Job not found"}), 404

        applications = Application.query.filter_by(job_id=job_id).all()
        if not applications:
            return jsonify({"message": "No applicants found for this job"}), 404

        parsed_cvs = ParsedCV.query.filter_by(job_id=job_id).all()
        parsed_cv_dict = {cv.candidate_id: cv for cv in parsed_cvs}

        # Prepare job description text
        job_text = ' '.join([
            str(job.education_requirement or ''),
            str(job.work_experience or ''),
            str(job.extra_skills or ''),
            str(job.job_requirements or ''),
            str(job.skills_needed or '')
        ])
        job_text_clean = clean_text(job_text)

        results = {
            'selected': [],
            'rejected': [],
            'errors': []
        }

        for app in applications:
            candidate_id = app.candidate_id
            parsed_cv = parsed_cv_dict.get(candidate_id)
            
            if not parsed_cv:
                results['errors'].append({
                    'candidate_id': candidate_id,
                    'error': 'No parsed CV data available'
                })
                continue

            # Get CV text from parsed data
            cv_text = ' '.join([
                ' '.join(parsed_cv.parsed_data.get(label.lower(), []))
                for labels in label_mapping.values()
                for label in labels
            ])
            cv_text_clean = clean_text(cv_text)

            # Calculate similarity score
            sbert_score = sbert_similarity_score(job_text_clean, cv_text_clean)
            
            # Determine selection status (65% threshold)
            if sbert_score >= 0.65:
                results['selected'].append(candidate_id)
            else:
                results['rejected'].append(candidate_id)

        return jsonify({
            "message": "Automatic evaluation completed",
            "results": results
        }), 200

    except Exception as e:
        logging.error(f"Error in auto-evaluation for job {job_id}: {e}")
        return jsonify({"message": f"An error occurred: {str(e)}"}), 500


@app.route('/hr/auto-notify-candidates/<int:job_id>', methods=['POST'])
def auto_notify_candidates(job_id):
    try:
        # First auto-evaluate candidates
        eval_response = auto_evaluate_candidates(job_id)
        if eval_response[1] != 200:
            return eval_response

        eval_data = eval_response[0].json
        results = eval_data.get('results', {})
        
        # Send notifications
        selected_success = []
        selected_failed = []
        for candidate_id in results.get('selected', []):
            try:
                parsed_cv = ParsedCV.query.filter_by(job_id=job_id, candidate_id=candidate_id).first()
                feedback = None
                if parsed_cv:
                    feedback_response = sbert_feedback(job_id, parsed_cv.id)
                    if feedback_response[1] == 200:
                        feedback = feedback_response[0].json.get('feedback_message', '')
                
                if send_candidate_notification(job_id, candidate_id, 'selected', feedback):
                    selected_success.append(candidate_id)
                else:
                    selected_failed.append(candidate_id)
            except Exception as e:
                logging.error(f"Error notifying selected candidate {candidate_id}: {e}")
                selected_failed.append(candidate_id)

        # Send rejections
        rejected_success = []
        rejected_failed = []
        for candidate_id in results.get('rejected', []):
            try:
                parsed_cv = ParsedCV.query.filter_by(job_id=job_id, candidate_id=candidate_id).first()
                feedback = None
                if parsed_cv:
                    feedback_response = sbert_feedback(job_id, parsed_cv.id)
                    if feedback_response[1] == 200:
                        feedback = feedback_response[0].json.get('feedback_message', '')
                
                if send_candidate_notification(job_id, candidate_id, 'rejected', feedback):
                    rejected_success.append(candidate_id)
                else:
                    rejected_failed.append(candidate_id)
            except Exception as e:
                logging.error(f"Error notifying rejected candidate {candidate_id}: {e}")
                rejected_failed.append(candidate_id)

        return jsonify({
            "message": "Automatic notification process completed",
            "results": {
                "selected_success": selected_success,
                "selected_failed": selected_failed,
                "rejected_success": rejected_success,
                "rejected_failed": rejected_failed,
                "errors": results.get('errors', [])
            }
        }), 200

    except Exception as e:
        logging.error(f"Error in auto-notification for job {job_id}: {e}")
        return jsonify({"message": f"An error occurred: {str(e)}"}), 500
    
@app.route('/hr/manual-notify-candidates/<int:job_id>', methods=['POST'])
def manual_notify_candidates(job_id):
    try:
        data = request.get_json()
        selected = data.get('selected', [])
        rejected = data.get('rejected', [])
        
        job = Job.query.get(job_id)
        if not job:
            return jsonify({"message": "Job not found"}), 404

        results = {
            'selected_success': [],
            'selected_failed': [],
            'rejected_success': [],
            'rejected_failed': [],
            'errors': []
        }
        
        # Notify selected candidates
        for candidate_id in selected:
            try:
                candidate = CandidateUser.query.get(candidate_id)
                if not candidate:
                    results['errors'].append(f"Candidate {candidate_id} not found")
                    continue

                parsed_cv = ParsedCV.query.filter_by(job_id=job_id, candidate_id=candidate_id).first()
                feedback = None
                
                if parsed_cv:
                    try:
                        # Get detailed feedback for the candidate
                        feedback_response = sbert_feedback(job_id, parsed_cv.id)
                        if feedback_response[1] == 200:
                            feedback = feedback_response[0].json.get('feedback_message', '')
                    except Exception as e:
                        logging.error(f"Error getting feedback for candidate {candidate_id}: {e}")
                
                # Send notification
                if send_candidate_notification(job_id, candidate_id, 'selected', feedback):
                    results['selected_success'].append(candidate_id)
                else:
                    results['selected_failed'].append(candidate_id)
                    
            except Exception as e:
                logging.error(f"Error processing selected candidate {candidate_id}: {e}")
                results['selected_failed'].append(candidate_id)
        
        # Notify rejected candidates
        for candidate_id in rejected:
            try:
                candidate = CandidateUser.query.get(candidate_id)
                if not candidate:
                    results['errors'].append(f"Candidate {candidate_id} not found")
                    continue

                parsed_cv = ParsedCV.query.filter_by(job_id=job_id, candidate_id=candidate_id).first()
                feedback = None
                
                if parsed_cv:
                    try:
                        # Get detailed feedback for the candidate
                        feedback_response = sbert_feedback(job_id, parsed_cv.id)
                        if feedback_response[1] == 200:
                            feedback = feedback_response[0].json.get('feedback_message', '')
                    except Exception as e:
                        logging.error(f"Error getting feedback for candidate {candidate_id}: {e}")
                
                # Send notification
                if send_candidate_notification(job_id, candidate_id, 'rejected', feedback):
                    results['rejected_success'].append(candidate_id)
                else:
                    results['rejected_failed'].append(candidate_id)
                    
            except Exception as e:
                logging.error(f"Error processing rejected candidate {candidate_id}: {e}")
                results['rejected_failed'].append(candidate_id)
        
        return jsonify({
            "message": "Manual notification process completed",
            "results": results
        }), 200
        
    except Exception as e:
        logging.error(f"Error in manual notification for job {job_id}: {e}")
        return jsonify({"message": f"An error occurred: {str(e)}"}), 500
from flask import Flask, Response
from fpdf import FPDF
import tempfile

@app.route('/feedback-pdf/<int:job_id>/<int:cv_id>', methods=['GET'])
def generate_feedback_pdf(job_id, cv_id):
    try:
        # Get data
        job = Job.query.get(job_id)
        parsed_cv = ParsedCV.query.get(cv_id)
        candidate = CandidateUser.query.get(parsed_cv.candidate_id)
        
        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        
        # Add a Unicode-compatible font (using built-in 'helvetica' as fallback)
        try:
            pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
            pdf.set_font('DejaVu', '', 12)
        except:
            pdf.set_font('helvetica', '', 12)  # Built-in font as fallback
        
        # Clean text function to handle special characters
        def clean_text(text):
            if not text:
                return ""
            # Replace problematic Unicode characters with ASCII equivalents
            replacements = {
                '’': "'",
                '‘': "'",
                '“': '"',
                '”': '"',
                '–': '-',
                '—': '-',
                '…': '...'
            }
            for k, v in replacements.items():
                text = text.replace(k, v)
            return text.encode('latin-1', 'ignore').decode('latin-1')
        
        # Header
        pdf.set_font('helvetica', 'B', 16)
        pdf.cell(200, 10, txt=clean_text("Feedback Report"), ln=1, align='C')
        pdf.set_font('helvetica', size=12)
        pdf.cell(200, 10, txt=clean_text(f"Candidate: {candidate.name}"), ln=1)
        pdf.cell(200, 10, txt=clean_text(f"Candidate Email: {candidate.email}"), ln=1)
        pdf.cell(200, 10, txt=clean_text(f"Position: {job.title} at {job.company_name}"), ln=1)
        pdf.ln(10)
        
        # Get feedback data
        feedback_data = sbert_feedback(job_id, cv_id)[0].json
        
        # Matching Items
        pdf.set_font('helvetica', 'B', 14)
        pdf.cell(200, 10, txt=clean_text("Strong Matches:"), ln=1)
        pdf.set_font('helvetica', size=12)
        
        for item in feedback_data.get('matching_items', []):
            pdf.cell(0, 10, txt=clean_text(f">> {item.get('field', '')}"), ln=1)
            pdf.multi_cell(0, 8, txt=clean_text(f"Job Requirement: {item.get('job_data', '')}"))
            pdf.multi_cell(0, 8, txt=clean_text(f"Your Skills: {item.get('cv_data', '')}"))
            pdf.ln(3)
        
        # Missing Items
        pdf.set_font('helvetica', 'B', 14)
        pdf.cell(200, 10, txt=clean_text("Areas Needing Improvement:"), ln=1)
        pdf.set_font('helvetica', size=12)
        
        for item in feedback_data.get('missing_items', []):
            pdf.cell(0, 10, txt=clean_text(f"!! {item.get('field', '')}"), ln=1)
            pdf.multi_cell(0, 8, txt=clean_text(f"Job Requirement: {item.get('job_data', '')}"))
            pdf.multi_cell(0, 8, txt=clean_text(f"Your Skills: {item.get('cv_data', 'Not specified')}"))
            pdf.ln(3)
        
        # Generate PDF
        pdf_output = pdf.output(dest='S').encode('latin-1', 'ignore')
        return Response(
            pdf_output,
            mimetype="application/pdf",
            headers={"Content-disposition": f"attachment; filename=feedback_{clean_text(candidate.name)}.pdf"}
        )
        
    except Exception as e:
        app.logger.error(f"PDF generation error: {str(e)}")
        return jsonify({"message": "Error generating PDF", "error": str(e)}), 500
@app.route('/api/generate-feedback-docx/<int:job_id>/<int:cv_id>', methods=['GET'])
def generate_feedback_docx(job_id, cv_id):
    try:
        job = Job.query.get(job_id)
        parsed_cv = ParsedCV.query.get(cv_id)
        candidate = CandidateUser.query.get(parsed_cv.candidate_id)

        if not job or not parsed_cv or not candidate:
            return jsonify({"message": "Job, CV or candidate not found"}), 404

        # Get feedback data
        feedback_response = sbert_feedback(job_id, cv_id)
        if feedback_response[1] != 200:
            return feedback_response
        
        feedback_data = feedback_response[0].json

        # Create DOCX
        doc = docx.Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run(f"Application Feedback for {job.title} at {job.company_name}")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = 1  # Center alignment
        
        # Candidate info
        doc.add_paragraph(f"Candidate: {candidate.name}", style='Heading 2')
        doc.add_paragraph(f"Email: {candidate.email}")
        doc.add_paragraph()
        
        # Overall score and feedback
        doc.add_paragraph(f"Overall Match Score: {feedback_data['sbert_score']*100:.2f}%", style='Heading 2')
        doc.add_paragraph(feedback_data['feedback_message'])
        doc.add_paragraph()
        
        # Matching items
        doc.add_paragraph("Matching Items:", style='Heading 2')
        
        for item in feedback_data['matching_items']:
            doc.add_paragraph(item['field'], style='Heading 3')
            doc.add_paragraph("Job Requirements:")
            doc.add_paragraph(item['job_data'])
            doc.add_paragraph("Your Experience:")
            doc.add_paragraph(item['cv_data'])
            doc.add_paragraph()
        
        # Missing items
        doc.add_paragraph("Areas for Improvement:", style='Heading 2')
        
        for item in feedback_data['missing_items']:
            doc.add_paragraph(item['field'], style='Heading 3')
            doc.add_paragraph("Job Requirements:")
            doc.add_paragraph(item['job_data'])
            doc.add_paragraph("Your Experience:")
            doc.add_paragraph(item['cv_data'] or "Not found")
            doc.add_paragraph()
        
        # Save to memory
        from io import BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Generate filename
        filename = f"feedback_{candidate.name.replace(' ', '_')}_{job.company_name.replace(' ', '_')}.docx"
        
        return Response(
            file_stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logging.error(f"Error generating DOCX feedback: {e}")
        return jsonify({"message": "Error generating feedback DOCX"}), 500
import os
from flask import send_from_directory

# Add this configuration at the top with other configs
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx'}

# Create uploads directory if it doesn't exist
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

@app.route('/hr/download-cv/<int:job_id>/<int:candidate_id>', methods=['GET'])
def download_cv(job_id, candidate_id):
    try:
        application = Application.query.filter_by(job_id=job_id, candidate_id=candidate_id).first()
        if not application:
            return jsonify({'message': 'CV not found'}), 404

        filepath = os.path.join(app.config['UPLOAD_FOLDER'], application.cv_filename)
        if not os.path.exists(filepath):
            return jsonify({'message': 'CV file not found on server'}), 404

        return send_file(filepath, as_attachment=True, download_name=application.cv_filename)
    except Exception as e:
        logging.error(f"Error downloading CV: {e}")
        return jsonify({'message': 'An error occurred while downloading the CV'}), 500

if __name__ == '__main__':
    app.run(debug=True)

